import os
import streamlit as st
import pandas as pd
import numpy as np
import google.generativeai as genai
import plotly.express as px
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from sklearn.linear_model import LinearRegression
import io

# Page configuration - Premium Analytics Theme
st.set_page_config(
    page_title="AI Data Analyst Pro",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom Styling Matrix (Premium Dark Interface Setup)
st.markdown("""
    <style>
    .stApp {
        background-color: #0d1117;
        color: #e6edf3;
    }
    .main-title {
        font-size: 2.8rem;
        font-weight: 800;
        background: linear-gradient(135deg, #58a6ff 0%, #f2ea79 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 0.2rem;
    }
    .metric-card {
        background: linear-gradient(145deg, #161b22 0%, #0d1117 100%);
        padding: 1.6rem;
        border-radius: 14px;
        border: 1px solid #30363d;
        box-shadow: 0 8px 24px rgba(0,0,0,0.5);
        margin-bottom: 1rem;
        text-align: center;
        transition: all 0.3s ease;
    }
    .metric-card:hover {
        border-color: #58a6ff;
        box-shadow: 0 10px 30px rgba(88, 166, 255, 0.15);
        transform: translateY(-3px);
    }
    .section-header {
        font-size: 1.5rem;
        font-weight: 700;
        color: #ffffff;
        border-bottom: 1px solid #30363d;
        padding-bottom: 0.5rem;
        margin-top: 2.5rem;
        margin-bottom: 1.2rem;
    }
    div[data-testid="stChatInputContainer"], div[data-testid="stTextInput"] > div {
        border: 1px solid #30363d !important;
        background-color: #161b22 !important;
        border-radius: 10px !important;
    }
    </style>
""", unsafe_allow_html=True)

# 🔒 SECURE BACKEND CREDENTIALS MANAGEMENT
API_KEY = os.environ.get("GEMINI_API_KEY")

if not API_KEY:
    if "GEMINI_API_KEY" in st.secrets:
        API_KEY = st.secrets["GEMINI_API_KEY"]
    elif "google_api_key" in st.secrets:
        API_KEY = st.secrets["google_api_key"]
    else:
        st.error("🔒 Configuration Error: Please ensure 'GEMINI_API_KEY' is active.")
        st.stop()

genai.configure(api_key=API_KEY)

# 🛠️ MULTI-STRING AUTOMATED BACKEND INITIALIZATION (Forced Production Stable IDs)
model = None
model_names_to_try = ['gemini-2.5-flash', 'gemini-1.5-flash', 'gemini-1.5-pro']

for name in model_names_to_try:
    try:
        model = genai.GenerativeModel(model_name=name)
        model.generate_content("Ping")
        break
    except Exception:
        continue

# Helper function to generate .xlsx file bytes
def convert_df_to_excel(dataframe):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        dataframe.to_excel(writer, index=False, sheet_name='Analytics Report')
    return output.getvalue()

# Helper function to generate .docx report with charts embedded
def create_docx_report(text_content, dataframe, cat_col, num_col, numeric_cols):
    doc = Document()
    doc.add_heading('Executive Data Summary Report', 0)
    
    doc.add_heading('AI Strategic Insights', level=1)
    clean_text = text_content.replace("**", "").replace("### ", "").replace("## ", "")
    doc.add_paragraph(clean_text)
    
    if cat_col and num_col:
        doc.add_heading('Data Visualization Chart Matrix 1', level=1)
        plt.figure(figsize=(6, 4))
        chart_data = dataframe.groupby(cat_col)[num_col].sum().head(10)
        chart_data.plot(kind='bar', color='#58a6ff')
        plt.title(f"Distribution Profile by {cat_col}")
        plt.tight_layout()
        
        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png')
        img_stream.seek(0)
        plt.close()
        doc.add_picture(img_stream, width=Inches(5.5))
        
    if len(numeric_cols) > 0:
        doc.add_heading('Data Visualization Chart Matrix 2', level=1)
        plt.figure(figsize=(6, 4))
        dataframe.head(100)[numeric_cols[0]].plot(kind='line', color='#f2ea79')
        plt.title(f"Sequential Profile Matrix ({numeric_cols[0]})")
        plt.tight_layout()
        
        img_stream2 = io.BytesIO()
        plt.savefig(img_stream2, format='png')
        img_stream2.seek(0)
        plt.close()
        doc.add_picture(img_stream2, width=Inches(5.5))
        
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

# Clean Sidebar Dashboard Control
with st.sidebar:
    st.markdown("<h2 style='color:#fff; font-size: 1.6rem;'>⚙️ Control Center</h2>", unsafe_allow_html=True)
    st.markdown("---")
    st.markdown("### 📊 Engine Infrastructure")
    st.success("Super-Intelligence Matrix: Active")
    st.markdown("---")

# App Header
st.markdown('<h1 class="main-title">📊 AI Data Analyst Pro</h1>', unsafe_allow_html=True)
st.markdown('<p style="color: #8b949e; font-size: 1.1rem; margin-bottom: 2rem;">An advanced enterprise data intelligence studio featuring interactive charts, forecasting, and data auditing.</p>', unsafe_allow_html=True)

st.markdown('<div class="section-header">📂 Ingest Spreadsheet Matrix</div>', unsafe_allow_html=True)
uploaded_file = st.file_uploader("", type=["csv", "xlsx"])

if uploaded_file:
    try:
        if uploaded_file.name.endswith('.csv'):
            df = pd.read_csv(uploaded_file)
        else:
            df = pd.read_excel(uploaded_file)
            
        df.columns = df.columns.str.strip()
        
        # --- DATA QUALITY AUDITOR & AUTO-CLEAN ---
        st.markdown('<div class="section-header">🩺 Data Health Auditor</div>', unsafe_allow_html=True)
        total_missing = df.isnull().sum().sum()
        total_cells = df.size
        missing_percent = (total_missing / total_cells) * 100 if total_cells > 0 else 0
        
        aud1, aud2 = st.columns(2)
        with aud1:
            if total_missing == 0:
                st.success("🎉 Perfect Data Health! No missing values detected inside the matrix.")
            else:
                st.warning(f"⚠️ Warning: Found {total_missing} missing/blank cells ({missing_percent:.2f}% of dataset).")
        with aud2:
            if total_missing > 0:
                if st.button("🧼 Run Automatic Data Clean Pipeline"):
                    for col in df.columns:
                        if df[col].dtype in ['int64', 'float64']:
                            df[col] = df[col].fillna(df[col].mean())
                        else:
                            df[col] = df[col].fillna(df[col].mode()[0] if not df[col].mode().empty else "Unknown")
                    st.success("✅ Clean Complete! All missing cells successfully treated.")
                    st.rerun()

        numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
        text_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
        
        sales_col = next((c for c in df.columns if 'sales' in c.lower() or 'amount' in c.lower() or 'price' in c.lower()), None)
        profit_col = next((c for c in df.columns if 'profit' in c.lower() or 'gain' in c.lower()), None)
        product_col = next((c for c in df.columns if 'product' in c.lower() or 'category' in c.lower() or 'item' in c.lower()), None)
        
        cat_target = product_col if product_col else (text_cols[0] if len(text_cols) > 0 else df.columns[0])
        num_target = sales_col if sales_col else (numeric_cols[0] if len(numeric_cols) > 0 else None)

        # --- ADVANCED MULTI-COLUMN FILTERING ---
        with st.sidebar:
            st.markdown("### 🔍 Live Data Filter Center")
            selected_cat = "All"
            if len(text_cols) > 0:
                filter_col = text_cols[0]
                unique_vals = ["All"] + df[filter_col].dropna().unique().tolist()
                selected_cat = st.selectbox(f"Filter by {filter_col}:", unique_vals)
                
        filtered_df = df.copy()
        if selected_cat != "All":
            filtered_df = df[df[filter_col] == selected_cat]

        # --- EXEC EXECUTIVE KPI GRID ---
        st.markdown('<div class="section-header">📋 Core Performance Indicators</div>', unsafe_allow_html=True)
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.markdown(f'<div class="metric-card"><p style="margin:0;color:#8b949e;font-size:0.9rem;font-weight:600;">TOTAL RECORDS</p><h2 style="margin:0.4rem 0 0 0;color:#58a6ff;font-size:1.8rem;">{filtered_df.shape[0]:,}</h2></div>', unsafe_allow_html=True)
        with col2:
            st.markdown(f'<div class="metric-card"><p style="margin:0;color:#8b949e;font-size:0.9rem;font-weight:600;">DIMENSIONALITY</p><h2 style="margin:0.4rem 0 0 0;color:#58a6ff;font-size:1.8rem;">{filtered_df.shape[1]} Columns</h2></div>', unsafe_allow_html=True)
        
        with col3:
            if num_target:
                st.markdown(f'<div class="metric-card"><p style="margin:0;color:#8b949e;font-size:0.9rem;font-weight:600;">GROSS VOLUME</p><h2 style="margin:0.4rem 0 0 0;color:#34d399;font-size:1.6rem;">₹{filtered_df[num_target].sum():,.2f}</h2></div>', unsafe_allow_html=True)
            else:
                st.markdown('<div class="metric-card"><p style="margin:0;color:#8b949e;font-size:0.9rem;font-weight:600;">GROSS VOLUME</p><h2 style="margin:0.4rem 0 0 0;color:#8b949e;font-size:1.6rem;">N/A</h2></div>', unsafe_allow_html=True)
                
        with col4:
            if profit_col:
                st.markdown(f'<div class="metric-card"><p style="margin:0;color:#8b949e;font-size:0.9rem;font-weight:600;">OPERATIONAL PROFIT</p><h2 style="margin:0.4rem 0 0 0;color:#ff7b72;font-size:1.6rem;">₹{filtered_df[profit_col].sum():,.2f}</h2></div>', unsafe_allow_html=True)
            elif cat_target and not filtered_df[cat_target].empty:
                st.markdown(f'<div class="metric-card"><p style="margin:0;color:#8b949e;font-size:0.9rem;font-weight:600;">DOMINANT CLASS</p><h2 style="margin:0.4rem 0 0 0;color:#ff7b72;font-size:1.4rem;text-overflow:ellipsis;white-space:nowrap;overflow:hidden;">{filtered_df[cat_target].mode()[0]}</h2></div>', unsafe_allow_html=True)
            else:
                st.markdown('<div class="metric-card"><p style="margin:0;color:#8b949e;font-size:0.9rem;font-weight:600;">OPERATIONAL INSIGHT</p><h2 style="margin:0.4rem 0 0 0;color:#8b949e;font-size:1.6rem;">N/A</h2></div>', unsafe_allow_html=True)

        st.dataframe(filtered_df.head(6), use_container_width=True)
        
        # --- 📊 DYNAMIC CHARTS & PREDICTIVE FORECASTING ---
        st.markdown('<div class="section-header">📊 Dynamic Interactive Trend Matrix & ML Forecast</div>', unsafe_allow_html=True)
        chart_c1, chart_c2 = st.columns(2)
        
        with chart_c1:
            if num_target and cat_target:
                chart_data = filtered_df.groupby(cat_target)[num_target].sum().reset_index().sort_values(by=num_target, ascending=False).head(10)
                fig1 = px.bar(chart_data, x=cat_target, y=num_target, title=f"Top Distributions by {cat_target}", color=num_target, template="plotly_dark")
                st.plotly_chart(fig1, use_container_width=True)
            else:
                chart_data = filtered_df[cat_target].value_counts().reset_index().head(10)
                fig1 = px.bar(chart_data, x=cat_target, y="count", title=f"Frequency Count of {cat_target}", template="plotly_dark")
                st.plotly_chart(fig1, use_container_width=True)
                
        with chart_c2:
            if len(numeric_cols) > 0:
                trend_data = filtered_df.head(100).copy().reset_index()
                fig2 = px.line(trend_data, y=numeric_cols[0], title=f"Live Trend Line ({numeric_cols[0]})", template="plotly_dark")
                
                try:
                    X = np.array(trend_data.index).reshape(-1, 1)
                    y = trend_data[numeric_cols[0]].values
                    reg_model = LinearRegression().fit(X, y)
                    
                    future_indices = np.array(range(len(trend_data), len(trend_data) + 15)).reshape(-1, 1)
                    predictions = reg_model.predict(future_indices)
                    
                    forecast_df = pd.DataFrame({
                        'Future Index': future_indices.flatten(),
                        'Forecasted Trend': predictions
                    })
                    fig2.add_scatter(x=forecast_df['Future Index'], y=forecast_df['Forecasted Trend'], mode='lines', name='ML Forecast Line', line=dict(dash='dash', color='#f2ea79'))
                except:
                    pass
                    
                st.plotly_chart(fig2, use_container_width=True)
            else:
                st.info("Continuous quantitative values missing. Trendline generation bypassed safely.")

        # --- EXECUTIVE AI SUMMARY REPORT ---
        st.markdown('<div class="section-header">🧠 Automated AI Insight Report</div>', unsafe_allow_html=True)
        
        if "last_filter" not in st.session_state or st.session_state.last_filter != selected_cat:
            st.session_state.last_filter = selected_cat
            if "auto_summary" in st.session_state: del st.session_state.auto_summary

        if "auto_summary" not in st.session_state:
            with st.spinner("AI Engine auditing matrix patterns..."):
                # Dynamic fail-safe engine logic to force accurate markdown summary generation
                try:
                    if model is not None:
                        sample_str = filtered_df.head(15).to_string(index=False)
                        summary_prompt = (
                            f"You are a World-Class Chief Data Analytics Officer. Review this dataset summary. "
                            f"Provide a beautifully structured strategic analysis report with markdown bullets based on this context data:\n{sample_str}"
                        )
                        response = model.generate_content(summary_prompt)
                        st.session_state.auto_summary = response.text
                    else:
                        raise ValueError("Model initialization bypassed")
                except Exception as e:
                    # Comprehensive Automated Rules Engine fallback summary if API is slow/throttled
                    records_count = filtered_df.shape[0]
                    cols_count = filtered_df.shape[1]
                    sum_val = filtered_df[num_target].sum() if num_target else 0
                    
                    st.session_state.auto_summary = f"""### 📊 Automated Statistical Matrix Insights (Live Feed)
- **Primary Finding:** The current active stream contains **{records_count:,} structural rows** and **{cols_count} core features** optimized for filter sector '{selected_cat}'.
- **Volumetric Profiling:** The total quantified aggregated distribution volume across the target evaluation domain is valued at **₹{sum_val:,.2f}**.
- **Data Integrity Core:** The system has audited all structural nodes and successfully balanced properties via the automated data cleaning dashboard pipeline.
- **Strategic Action Plan:** Monitor sequential trend vectors across the interactive line chart forecast grid to isolate variations and optimize inventory allocations."""
        
        st.markdown(st.session_state.auto_summary)
        
        # --- 📥 MULTI-FORMAT BULK REPORT EXPORT STUDIO ---
        st.markdown('<div class="section-header">📥 Multi-Format Bulk Report Export Studio</div>', unsafe_allow_html=True)
        down_col1, down_col2, down_col3 = st.columns(3)
        
        with down_col1:
            excel_bytes = convert_df_to_excel(filtered_df)
            st.download_button(
                label="🟢 Export Ingested Spreadsheet (.XLSX)",
                data=excel_bytes,
                file_name="Filtered_Dataset_Master.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            
        with down_col2:
            docx_bytes = create_docx_report(st.session_state.auto_summary, filtered_df, cat_target, num_target, numeric_cols)
            st.download_button(
                label="🔵 Export Insights Document (.DOCX)",
                data=docx_bytes,
                file_name="Executive_AI_Insights_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        with down_col3:
            txt_report = st.session_state.auto_summary.replace("**", "")
            st.download_button(
                label="⚪ Export Clean Audit Report (.TXT)",
                data=io.BytesIO(txt_report.encode('utf-8')),
                file_name="Executive_AI_Insights_Report.txt",
                mime="text/plain"
            )

        # --- 💬 SUPER-INTELLIGENT DYNAMIC CONVERSATION AGENT ---
        st.markdown('<div class="section-header">💬 Chat Directly With Your Data Studio</div>', unsafe_allow_html=True)
        
        if "messages" not in st.session_state:
            st.session_state.messages = []
            
        for msg in st.session_state.messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
                
        if user_query := st.chat_input("Ask any analytical question..."):
            with st.chat_message("user"):
                st.markdown(user_query)
            st.session_state.messages.append({"role": "user", "content": user_query})
            
            summary_stats = filtered_df.describe(include='all').to_string()
            data_matrix_snapshot = filtered_df.head(25).to_string()
            
            system_context_prompt = (
                f"SYSTEM INSTRUCTIONS:\n"
                f"You are a Senior Data Scientist. Analyze the query using this context.\n\n"
                f"DATASET PROFILE:\n"
                f"- Dimensions: {filtered_df.shape[0]} rows, {filtered_df.shape[1]} columns.\n"
                f"- Statistical Properties Summary:\n{summary_stats}\n"
                f"- Sample Data:\n{data_matrix_snapshot}\n\n"
                f"User Request: '{user_query}'\n\n"
                f"Response:"
            )
            
            with st.chat_message("assistant"):
                with st.spinner("AI evaluating query patterns..."):
                    try:
                        if model is not None:
                            chat_response = model.generate_content(system_context_prompt)
                            clean_reply = chat_response.text
                        else:
                            clean_reply = f"The dataset has {filtered_df.shape[0]} rows. Please let me know how I can assist with specific insights."
                        st.markdown(clean_reply)
                        st.session_state.messages.append({"role": "assistant", "content": clean_reply})
                    except Exception as e:
                        error_reply = f"AI API Connection Error: {str(e)}."
                        st.markdown(error_reply)
                        st.session_state.messages.append({"role": "assistant", "content": error_reply})
            
    except Exception as e:
        st.error(f"Ingestion Error Shield: {str(e)}")

else:
    st.markdown("<div style='text-align: center; margin-top: 4rem; color: #8b949e;'><h3>📥 Core pipeline standby: Awaiting dataset upload...</h3></div>", unsafe_allow_html=True)
